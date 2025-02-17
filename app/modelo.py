from docx import Document
from docx.shared import Pt
from app import database_sqlite
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import os

class curso_unitario:
    def __init__(self, cursos):  
        self.instituicao = cursos[1]
        self.curso = cursos[2]
        self.conclusao = cursos[3]
        self.list = [
            f"▷ {self.instituicao}",
            [
                f"• {self.curso}",
                f"• Conclusão em {self.conclusao}"
            ]
        ]
        pass

class experiencia_unitario:
    def __init__(self, experiencias):
        self.tempo_inicio = experiencias[1]
        self.tempo_fim = experiencias[2]
        self.empresa = experiencias[3]
        self.cargo = experiencias[4]
        self.list = [
          f"{self.tempo_inicio} - {self.tempo_fim}",
            [
                f"{self.empresa}",
                f"{self.cargo}"
            ]  
        ]

class ProcessarModelo:
    def __init__(self, id):
        self.db = database_sqlite.DatabaseSqlite()
        self.curr_info = self.db.get_curriculo_by_id(id)
        self.cursos = self.db.get_cursos_by_curriculo_id(id)
        self.experiencias = self.db.get_experiencias_by_curriculo_id(id)
        self.processar()
        pass

    def processar(self):
        if not self.curr_info:
            return False
        
        cursos = []
        experiencias = []

        for curso in self.cursos:
            cursos.append(curso_unitario(curso).list)
        for experiencia in self.experiencias:
            experiencias.append(experiencia_unitario(experiencia).list)

        substituicoes = {
            "##NOME##": self.curr_info[1],
            "##ENDERECO##": self.curr_info[2],
            "##TELEFONE##": self.curr_info[3],
            "##EMAIL##": self.curr_info[4],
            "##OBJETIVO##": self.curr_info[5],
            "##PERFIL##": self.curr_info[6],
            "##ESCOLARIDADE##": None,
            "##EXPERIENCIA##": None
        }
        doc = Document("modelo.docx")
        for para in doc.paragraphs:
            for run in para.runs:
                for chave, valor in substituicoes.items():
                    if chave in run.text:
                        if chave == "##NOME##":
                            run.text = run.text.replace(chave, str(valor).upper())
                            run.bold = True
                            run.font.size = Pt(21)
                        if chave == "##ENDERECO##":
                            run.text = run.text.replace(chave, str(valor).capitalize())
                            run.font.size = Pt(12)
                        if chave == "##TELEFONE##":
                            run.text = run.text.replace(chave, str(valor))
                            run.font.size = Pt(12)
                        if chave == "##EMAIL##":
                            run.text = run.text.replace(chave, str(valor))
                            run.font.size = Pt(12)
                        if chave == "##OBJETIVO##":
                            run.text = run.text.replace(chave, str(valor).capitalize())
                            run.font.size = Pt(12)
                            run._element.getparent().alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        if chave == "##PERFIL##":
                            run.text = run.text.replace(chave, str(valor).capitalize())
                            run.font.size = Pt(12)
                            run._element.getparent().alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                        if chave == "##ESCOLARIDADE##":
                            run.text = run.text.replace(chave, '')  
                         
                            for curso, subitens in cursos:
                              
                                run_titulo = para.add_run(f"\n{curso.upper()} ")  
                                run_titulo.bold = True  
                                run_titulo.font.size = Pt(14)  

                                for subitem in subitens:
                                    run_subitem = para.add_run(f"\n    {subitem}") 
                                    run_subitem.font.size = Pt(12)  
                                    run_subitem.bold = False  
                                    run_subitem.italic = True
                        if chave == "##EXPERIENCIA##":
                            run.text = run.text.replace(chave, '')  
                         
                            for experiencia, subitens in experiencias:
                              
                                run_titulo = para.add_run(f"\n{experiencia} ")  
                                run_titulo.bold = True
                                run_titulo.italic = True  
                                run_titulo.font.size = Pt(14)  

                                for subitem in subitens:
                                    run_subitem = para.add_run(f"\n{subitem}") 
                                    run_subitem.font.size = Pt(12)  
                                    run_subitem.bold = False
        
        folder_path = 'app\curriculo'
        os.makedirs(folder_path, exist_ok=True)  

        doc.save(f"{folder_path}/curriculo[{self.curr_info[0]}].docx")
        return True


